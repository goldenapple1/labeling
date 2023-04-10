import typing
import time
import os
import spacy
import pymysql
from docx import Document
from datetime import datetime

start = time.time()
collins = {5: -1, 4: -1, 3: 0, 2: 1, 1: 2, 0: 3}
oxford = {1: -3, 0: 0}
tag = {"zk": -2, "gk": -1, "cet4": 0, "ky": 1, "cet6": 1, "toefl": 1, 
       "ielts": -1, "gre": 1}
maximum =[-1, -1, 0, 0, 1]
catagory = ["spoken language", "fiction", "megazine", "newspaper", 
            "academic material"]
endnotes = {}
label_index = []
log = open("log.txt", "w+", encoding="utf-8")
secrets = open("secrets.txt", "r", encoding="utf-8").readlines()

path = secrets[0][:-1]
pwd = secrets[1]

docx_files = []
latest_date = None
latest_file = None
now = datetime.now()
month = str(now.month)

if len(month) == 1:
    month = "0"+month
date = str(now.year) + "-" + month

for n in os.listdir(path+"\\"+date):
    if n[-5:] == ".docx":
        file_date = os.path.getmtime(path+"\\"+date+"\\"+n)

        if latest_date is None or file_date > latest_date:
            latest_date = file_date
            latest_file = n

doc = Document(path+"\\"+date+"\\"+latest_file)

spacy.require_gpu()
nlp = spacy.load('en_core_web_trf')

db = pymysql.connect(
    host="localhost", 
    user="root", 
    password=pwd, 
    database="dict"
)

cursor = db.cursor()


def find_max(l:list) -> int:
    local_max = 0
    idx = 0
    i = 0

    for n in l:
        if n > local_max:
            local_max = n
            idx = i
        i += 1
    
    return idx


def evaluation(data:tuple) -> str:
    cursor.execute(f'SELECT * FROM freq WHERE word="{data[1]}";')
    freq = cursor.fetchone()
    score = 5
    
    score += collins[data[2]]
    score += oxford[data[3]]

    if data[4] == "":
        score += 1
    else:
        counter = 0
        for n in data[4].split(" "):
            score += tag[n]
            counter += 1

        if counter > 3:
            score -= 1
    
    if freq != None:
        highest = find_max(freq[4:])
        score += maximum[highest]

        if freq[0] <= 4500:
            score -= 2
        elif freq[0] <= 7500:
            score -= 1
        elif freq[0] > 9000:
            score += 1

        most_freq = f", appear most in {catagory[highest]}"
    else:
        most_freq = ""

    if score > 7:
        return f"score: {score}{most_freq}"
    return ""


def word_query(word:str) -> list:
    if word in endnotes:
        return []

    cursor.execute(f'SELECT * FROM words WHERE word="{word}";')
    word_data = cursor.fetchone()
    if word_data == None:
        return []

    score = evaluation(word_data)
    anno = word

    if word_data[5] != "":
        if word_data[5] in endnotes:
            return[]
        
        cursor.execute(f'SELECT * FROM words WHERE word="{word_data[5]}";')
        lemma_data = cursor.fetchone()

        if lemma_data != None:
            lemma_score = evaluation(lemma_data)

            if lemma_score != "" and score != "":
                if int(lemma_score.split(",")[0][7:]) < int(score.split(",")[0][7:]):
                    score = lemma_score
                    anno = word_data[5]

            elif lemma_score == "":
                return []
    
    if score != "":
        cursor.execute(f'SELECT * FROM dict WHERE word="{anno}";')
        meaning = cursor.fetchone()

        if meaning[2] == "":
            pos2 = ""
        else:
            pos2 = f"/{meaning[2]}/"

        if meaning != None:
            return [word, f"{pos2}: {meaning[3]}\n{meaning[4]} -{score}"]
        
    return []


def tokenize(text):
    entities = nlp(text).ents
    idx = {ent.start_char: ent.end_char for ent in entities}
    labels = []
    temp: str = ""
    index = 0
    ignore = -1

    while index < len(text):
        
        if index in idx:
            ignore = idx[index]
            index += 1
            continue

        elif ignore != -1:
            if index != ignore:
                index += 1
                continue
            else:
                ignore = -1

        if text[index].isalpha():
            temp += text[index]

        elif text[index] in "'’-":
            temp += text[index]

        elif temp != "":
            if len(temp) > 2 and "’" not in temp:
                result = word_query(temp.lower())

                if result != [] and result[0] not in endnotes:
                    endnotes[result[0]] = result[1]
                    print(f"{result[0]}", file=log)
                    labels.append(index)
                
            temp = ""

        index += 1
    
    label_index.append(labels)


num_para = len(doc.paragraphs)
mark_count = 1
for i in range(num_para):
    para_text = doc.paragraphs[i].text
    tokenize(para_text)

    if label_index[i] != []:
        new_text = list(para_text)
        new_text.insert(label_index[i][0], f"({mark_count}.)")
        mark_count += 1

        p = 1
        while p < len(label_index[i]):
            new_text.insert(label_index[i][p] + (p - 1) + 1, 
                            f"({mark_count}.)")
            p += 1
            mark_count += 1

        doc.paragraphs[i].text = new_text

count = 1
for n in endnotes:
    new_para = doc.add_paragraph(f'({count}.) ')
    new_para.add_run(n).bold = True
    new_para.add_run(endnotes[n])
    count += 1

doc.save(f'{latest_file[:-5]}_annotated.docx')
print(f"time spent: {(time.time() - start):.2f}")